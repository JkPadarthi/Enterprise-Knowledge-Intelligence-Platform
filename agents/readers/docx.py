"""DOCX reader (optional: requires ``python-docx``).

The optional dependency is imported lazily so the package imports cleanly even
when ``python-docx`` is not installed; :meth:`read` raises a clear ``RuntimeError``
explaining how to enable ``.docx`` support.
"""

from __future__ import annotations

from agents.readers.base import BaseReader, ReaderResult

try:
    from docx import Document as _DocxDocument  # type: ignore

    _DOCX_AVAILABLE = True
except Exception:  # noqa: BLE001 - optional dependency, import is best-effort
    _DOCX_AVAILABLE = False


class DocxReader(BaseReader):
    """Extracts paragraph and table-cell text from a ``.docx`` file."""

    extensions = (".docx",)

    def read(self, data: bytes, filename: str = "") -> ReaderResult:
        if not _DOCX_AVAILABLE:
            raise RuntimeError(
                "DocxReader requires the optional 'python-docx' package. "
                "Install it with: pip install python-docx"
            )

        import io

        document = _DocxDocument(io.BytesIO(data))
        paragraphs: list[str] = [
            p.text for p in document.paragraphs if p.text and p.text.strip()
        ]
        # Pull table cell text so structured content is not dropped.
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        paragraphs.append(cell.text)

        text = "\n".join(paragraphs)
        return ReaderResult(raw_text=text, pages=[text], num_pages=1, format="docx")
